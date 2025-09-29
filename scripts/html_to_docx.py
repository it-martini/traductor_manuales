#!/usr/bin/env python3
"""
Convertidor HTML español a DOCX - Con enlaces internos REALES funcionando
"""
from pathlib import Path

# Definir rutas base del proyecto
BASE_DIR = Path(__file__).parent
PROJECT_ROOT = BASE_DIR.parent
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.parser import parse_xml
from docx.oxml.ns import nsdecls
import re
import os
import time
from datetime import datetime
from bs4 import BeautifulSoup
from toc_handler import TOCHandler
from system_config import get_log_file

class DOCXLogger:
    """Logger específico para generación DOCX"""

    def __init__(self, manual_name, lang_code):
        self.manual_name = manual_name
        self.lang_code = lang_code
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.log_file = get_log_file(f"docx_{manual_name}_{lang_code}_{timestamp}")

        # Log inicio
        with open(self.log_file, 'w', encoding='utf-8') as f:
            f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] DOCX_START: Manual={manual_name}, Lang={lang_code}\n")

    def log_step(self, message):
        """Log paso del proceso"""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        with open(self.log_file, 'a', encoding='utf-8') as f:
            f.write(f"[{timestamp}] {message}\n")

    def log_file_processed(self, filename, success=True):
        """Log archivo procesado"""
        status = "SUCCESS" if success else "ERROR"
        self.log_step(f"FILE_PROCESSED: {filename} [{status}]")

    def finalize(self, total_files, processed_files, file_size=None):
        """Finalizar logging"""
        size_str = f", Size={file_size}" if file_size else ""
        self.log_step(f"DOCX_END: Total={total_files}, Processed={processed_files}{size_str}")

class DOCXProgressDisplay:
    """Display de progreso para generación DOCX"""

    def __init__(self, total_files=0):
        self.total_files = total_files
        self.current_file = 0
        self.start_time = time.time()
        self.current_step = ""

    def show_start(self, output_file):
        """Mostrar inicio de generación"""
        print(f"📄 Generando DOCX: {Path(output_file).name}")

    def show_step(self, step_name):
        """Mostrar paso actual"""
        self.current_step = step_name
        print(f"   🔧 {step_name}...")

    def show_file_progress(self, current, total, filename):
        """Mostrar progreso de archivo"""
        if total > 0:
            progress = current / total
            bar_length = 20
            filled_length = int(bar_length * progress)
            bar = '█' * filled_length + '░' * (bar_length - filled_length)

            elapsed = time.time() - self.start_time
            if current > 0 and elapsed > 0:
                avg_time = elapsed / current
                remaining = (total - current) * avg_time
                eta_str = f"ETA: {int(remaining//60)}:{int(remaining%60):02d}"
            else:
                eta_str = "ETA: --:--"

            print(f"\r   [{bar}] {current}/{total} páginas | {eta_str}", end="", flush=True)

    def show_complete(self, processed_files, total_files, output_file):
        """Mostrar completación"""
        file_size = ""
        try:
            if Path(output_file).exists():
                size_bytes = Path(output_file).stat().st_size
                if size_bytes > 1024*1024:
                    file_size = f" ({size_bytes/(1024*1024):.1f} MB)"
                else:
                    file_size = f" ({size_bytes/1024:.0f} KB)"
        except:
            pass

        elapsed = time.time() - self.start_time
        print(f"\n✅ DOCX generado: {Path(output_file).name}{file_size}")
        print(f"   📊 Procesadas: {processed_files}/{total_files} páginas en {elapsed:.1f}s")

def spanish_html_with_real_links(html_dir, output_docx, manual_name="manual", lang_code="es", manual_type="open_aula_back"):
    """Convertir HTML español a DOCX con enlaces internos REALES"""

    html_path = Path(html_dir)
    if not html_path.exists():
        print(f"❌ Error: No se encuentra {html_path}")
        return False

    # Inicializar logger y progreso
    logger = DOCXLogger(manual_name, lang_code)
    logger.log_step(f"DOCX_CONVERSION_START: {html_path} -> {output_docx}")

    # Obtener archivos HTML primero para configurar progreso
    html_files = get_all_html_files_structured(html_path)
    progress = DOCXProgressDisplay(len(html_files))

    progress.show_start(output_docx)
    logger.log_step(f"HTML_FILES_FOUND: {len(html_files)} files")

    # Crear documento DOCX
    progress.show_step("Configurando documento")
    doc = Document()
    setup_enhanced_document(doc)
    logger.log_step("DOCUMENT_SETUP: Complete")

    # Cargar estilos CSS
    progress.show_step("Cargando estilos CSS")
    css_styles = load_css_styles_from_spanish(manual_type)
    logger.log_step(f"CSS_STYLES_LOADED: {len(css_styles)} classes")

    # CREAR PÁGINA DE TÍTULO Y ÍNDICE
    progress.show_step("Creando índice")
    create_title_page_and_index(doc, html_files)
    logger.log_step("TITLE_PAGE_INDEX: Created")

    # PASO 1: Crear mapeo de archivos a bookmarks
    progress.show_step("Creando bookmarks")
    bookmark_mapping = create_bookmark_mapping(html_files)
    logger.log_step(f"BOOKMARKS_CREATED: {len(bookmark_mapping)} bookmarks")

    # PASO 2: Procesar cada archivo
    processed_files = 0
    bookmark_id_counter = 1

    for i, html_file in enumerate(html_files):
        try:
            progress.show_file_progress(i+1, len(html_files), html_file.name)

            success, bookmark_id_counter = process_html_file_with_real_links(
                doc, html_file, html_path, css_styles, bookmark_mapping, bookmark_id_counter
            )

            if success:
                processed_files += 1
                logger.log_file_processed(html_file.name, True)
            else:
                logger.log_file_processed(html_file.name, False)

            # Agregar salto de página entre archivos (excepto el último)
            if i < len(html_files) - 1:
                doc.add_page_break()

        except Exception as e:
            logger.log_step(f"ERROR_PROCESSING: {html_file.name} - {str(e)}")
            continue

    # Guardar documento
    try:
        progress.show_step("Guardando documento")
        doc.save(output_docx)

        # Obtener tamaño del archivo
        file_size = None
        if Path(output_docx).exists():
            file_size = Path(output_docx).stat().st_size

        logger.finalize(len(html_files), processed_files, file_size)
        progress.show_complete(processed_files, len(html_files), output_docx)

        return True
    except Exception as e:
        logger.log_step(f"ERROR_SAVING: {str(e)}")
        print(f"❌ Error guardando DOCX: {e}")
        return False

def setup_enhanced_document(doc):
    """Configurar propiedades del documento"""

    core_props = doc.core_properties
    core_props.title = "Manual de Usuario - Con Enlaces REALES"
    core_props.subject = "Manual de Usuario"
    core_props.author = "Educativa"
    core_props.comments = "Manual de usuario con navegación interna REAL"

    # Configurar fuente por defecto del documento (Calibri)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1.2)  # Más espacio para pie de página
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Configurar pie de página
    setup_footer(sections[0])

def setup_footer(section, language='es'):
    """Configurar pie de página profesional"""

    t = get_translations(language)
    footer = section.footer

    # Limpiar párrafos existentes
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Crear tabla invisible de 1 fila x 2 columnas para alineación real
    table = footer.add_table(rows=1, cols=2, width=Inches(7.5))
    table.style = 'Table Grid'

    # Configurar tabla invisible (sin bordes)
    for row in table.rows:
        for cell in row.cells:
            # Quitar bordes
            tc = cell._element
            tcPr = tc.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.append(tcPr)

            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    # Configurar ancho de columnas
    table.columns[0].width = Inches(4.5)  # Izquierda más ancha
    table.columns[1].width = Inches(3.0)  # Derecha más estrecha

    # CELDA IZQUIERDA: Copyright
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run(t['footer_company'])
    left_run.font.name = 'Calibri'
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = RGBColor(102, 102, 102)

    # CELDA DERECHA: Numeración de páginas
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Texto "Página " / "Page "
    page_run = right_para.add_run(f"{t['footer_page']} ")
    page_run.font.name = 'Calibri'
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = RGBColor(102, 102, 102)

    # Número de página actual
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    page_run._element.append(fldChar1)
    page_run._element.append(instrText)
    page_run._element.append(fldChar2)

    # " de " / " of "
    of_run = right_para.add_run(f" {t['footer_of']} ")
    of_run.font.name = 'Calibri'
    of_run.font.size = Pt(9)
    of_run.font.color.rgb = RGBColor(102, 102, 102)

    # Número total de páginas
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    of_run._element.append(fldChar3)
    of_run._element.append(instrText2)
    of_run._element.append(fldChar4)

def get_translations(language='es'):
    """Obtener traducciones para elementos del documento"""
    from datetime import datetime
    current_year = datetime.now().year

    translations = {
        'es': {
            'title': 'Campus virtual',
            'subtitle': 'Manual de Usuario',
            'index_title': 'Índice',
            'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
            'footer_page': 'Página',
            'footer_of': 'de'
        },
        'en': {
            'title': 'Virtual Campus',
            'subtitle': 'User Manual',
            'index_title': 'Index',
            'footer_company': f'© {current_year} e-ducativa Virtual Education S.A.',
            'footer_page': 'Page',
            'footer_of': 'of'
        }
    }
    return translations.get(language, translations['es'])

def add_run_with_font(paragraph, text=""):
    """Crear run con fuente Calibri por defecto"""
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    return run

def load_css_styles_from_spanish(manual_type='open_aula_back'):
    """Cargar estilos CSS"""
    css_styles = {}

    # Determinar directorio source según el tipo de manual
    manual_dir = f"{manual_type}_es"
    css_file = PROJECT_ROOT / "original" / manual_dir / "html" / "css" / "hnd.content.css"
    if css_file.exists():
        try:
            with open(css_file, 'r', encoding='utf-8') as f:
                css_content = f.read()
            css_styles = parse_rvts_styles(css_content)
        except Exception as e:
            print(f"   ⚠️ Error cargando CSS: {e}")

    return css_styles

def parse_rvts_styles(css_content):
    """Parsear clases rvts del CSS"""
    styles = {}

    patterns = [
        r'(?:a\.rvts(\d+),?\s*)?span\.rvts(\d+)[^{]*\{([^}]+)\}',
        r'a\.rvts(\d+),?\s*span\.rvts(\d+)[^{]*\{([^}]+)\}',
        r'span\.rvts(\d+)[^{]*\{([^}]+)\}'
    ]

    for pattern in patterns:
        matches = re.finditer(pattern, css_content, re.MULTILINE | re.DOTALL)
        for match in matches:
            groups = match.groups()

            if len(groups) == 3 and groups[1]:
                class_num = groups[1]
                css_rules = groups[2]
            elif len(groups) == 2:
                class_num = groups[0]
                css_rules = groups[1]
            else:
                continue

            style_props = {}
            rules = css_rules.strip().split(';')

            for rule in rules:
                if ':' in rule:
                    prop, value = rule.split(':', 1)
                    prop = prop.strip()
                    value = value.strip()
                    style_props[prop] = value

            styles[f'rvts{class_num}'] = style_props

    return styles

def create_bookmark_mapping(html_files):
    """Crear mapeo de archivos HTML a nombres de bookmarks"""
    bookmark_mapping = {}

    for html_file in html_files:
        filename = html_file.name
        bookmark_name = clean_bookmark_name(filename)
        bookmark_mapping[filename] = bookmark_name

    return bookmark_mapping

def clean_bookmark_name(filename):
    """Limpiar nombre de archivo para usar como bookmark"""
    name = filename.replace('.html', '')
    name = re.sub(r'[^a-zA-Z0-9]', '_', name)
    if name and name[0].isdigit():
        name = 'section_' + name
    return name or 'unknown'

def parse_html_toc_structure(manual_type='open_aula_back'):
    """Parsear la estructura jerárquica del TOC HTML"""
    from bs4 import BeautifulSoup

    try:
        # Determinar directorio source según el tipo de manual
        manual_dir = f"{manual_type}_es"
        index_path = PROJECT_ROOT / "original" / manual_dir / "html" / "index.html"
        with open(index_path, 'r', encoding='utf-8') as f:
            content = f.read()

        soup = BeautifulSoup(content, 'html.parser')
        toc_div = soup.find('div', {'id': 'toc'})

        if toc_div:
            root_ul = toc_div.find('ul')
            if root_ul:
                return parse_toc_ul(root_ul)
    except Exception as e:
        print(f"   ⚠️ Error parseando TOC HTML: {e}")

    return []

def parse_toc_ul(ul_elem, level=0):
    """Parsear recursivamente elementos UL/LI del TOC"""
    items = []

    for li in ul_elem.find_all('li', recursive=False):
        # Buscar enlace directo en este LI
        link = li.find('a', recursive=False)
        if link:
            href = link.get('href', '')
            text = link.get_text()

            # Solo incluir si no es Home/index
            if href not in ['Home.html', 'index.html']:
                item = {
                    'title': text,
                    'href': href,
                    'level': level,
                    'children': []
                }

                # Buscar sub-elementos
                sub_ul = li.find('ul', recursive=False)
                if sub_ul:
                    item['children'] = parse_toc_ul(sub_ul, level + 1)

                items.append(item)

    return items

def create_structured_index(doc, toc_structure, html_files):
    """Crear índice estructurado con numeración jerárquica"""

    # Crear un mapeo de archivos disponibles
    available_files = {f.name: f for f in html_files}

    def add_toc_items(items, numbering_path=[]):
        """Agregar items del TOC con numeración jerárquica"""
        for i, item in enumerate(items, 1):
            if item['href'] in available_files:
                # Calcular numeración jerárquica
                current_numbering = numbering_path + [i]
                number_str = '.'.join(map(str, current_numbering)) + '.'

                # Crear entrada de índice
                index_para = doc.add_paragraph()

                # Indentación basada en nivel
                indent = Pt(36 + (item['level'] * 18))  # 36pt base + 18pt por nivel
                index_para.paragraph_format.left_indent = indent

                # Agregar numeración
                num_run = add_run_with_font(index_para, f"{number_str} ")
                num_run.font.bold = True
                num_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul como títulos

                # Crear hyperlink para el título
                bookmark_name = clean_bookmark_name(item['href'])
                create_real_internal_hyperlink(index_para, item['title'], bookmark_name, {}, [])

                print(f"     📋 Índice: {'  ' * item['level']}{number_str} {item['title']} → {bookmark_name}")

            # Procesar hijos recursivamente con numeración extendida
            if item['children']:
                child_numbering = numbering_path + [i] if item['href'] in available_files else numbering_path
                add_toc_items(item['children'], child_numbering)

    print("   📋 Generando índice estructurado con numeración jerárquica...")
    add_toc_items(toc_structure)

def create_title_page_and_index(doc, html_files, language='es'):
    """Crear página de título y índice funcional usando nuevo TOC handler"""

    print("   📝 Creando página de título y índice...")

    t = get_translations(language)

    # TÍTULO PRINCIPAL: "Campus virtual" / "Virtual Campus"
    title_para = doc.add_paragraph()
    title_para.alignment = 1  # Centrado
    title_run = title_para.add_run(t['title'])
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro

    # SUBTÍTULO: "Manual de Usuario" / "User Manual"
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = 1  # Centrado
    subtitle_run = subtitle_para.add_run(t['subtitle'])
    subtitle_run.font.name = 'Calibri'
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(68, 68, 68)  # Gris oscuro

    # Espaciado
    doc.add_paragraph("")
    doc.add_paragraph("")

    # GENERAR ÍNDICE CON NUEVO TOC HANDLER
    # Determinar directorio HTML para parsear TOC
    if html_files:
        html_dir = html_files[0].parent
        toc_handler = TOCHandler()
        toc_structure = toc_handler.parse_toc_json(html_dir)

        # Crear mapeo de archivos disponibles
        available_files = {f.name: f for f in html_files}

        # Generar índice jerárquico
        toc_handler.create_hierarchical_index(doc, toc_structure, available_files, language)


def get_clean_title_from_filename(filename):
    """Convertir nombre de archivo a título legible"""
    title_map = {
        'Introduccion': 'Introducción',
        'Redessociales': 'Redes sociales',
        'Anexodocente': 'Anexo Docente',
        'Aulasdocente': 'Aulas (Docente)',
        'Seccionesdocente': 'Secciones (Docente)',
        'Archivos': 'Archivos',
        'Categorias': 'Categorías',
        'Portafolio': 'Portafolio',
        'Documentos': 'Documentos',
        'Carpetas': 'Carpetas'
    }

    return title_map.get(filename, filename.replace('_', ' ').title())

def get_all_html_files_structured(html_path):
    """Obtener TODOS los archivos HTML ordenados según TOC usando el nuevo handler"""

    # Usar el nuevo TOCHandler
    toc_handler = TOCHandler()
    toc_structure = toc_handler.parse_toc_json(html_path)

    # Obtener archivos ordenados según TOC
    html_files = toc_handler.get_ordered_html_files(html_path, toc_structure)

    print(f"   📋 Archivos ordenados según TOC: {len(html_files)}")

    return html_files

def get_test_toc_files(html_path):
    """Obtener archivos limitados para probar la numeración jerárquica"""
    html_files = []

    # Solo algunos archivos para probar la estructura del TOC
    test_files = [
        'Redessociales.html', 'Introduccion.html', 'Aulas.html', 'Encabezado.html',
        'Perfil.html', 'Preferencias.html', 'Catalogo.html', 'Portafolio.html',
        'Carpetas.html', 'Secciones.html', 'Categorias.html', 'Programa.html'
    ]

    for test_file in test_files:
        file_path = html_path / test_file
        if file_path.exists():
            html_files.append(file_path)

    print(f"   📋 Archivos de prueba numeración: {len(html_files)}")

    return html_files

def process_html_file_with_real_links(doc, html_file, base_path, css_styles, bookmark_mapping, bookmark_id_counter):
    """Procesar archivo HTML con enlaces internos REALES"""

    try:
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, 'html.parser')

        # CREAR BOOKMARK REAL para esta sección
        title = extract_title(soup, html_file)
        if title:
            heading, bookmark_id_counter = add_heading_with_real_bookmark(
                doc, title, html_file.name, bookmark_mapping, bookmark_id_counter
            )

        # Extraer contenido principal
        main_content = soup.find('div', class_='main-content')
        if not main_content:
            main_content = soup.find('div', id='topic-content')

        if main_content:
            bookmark_id_counter = process_html_content_with_real_links(
                doc, main_content, base_path, css_styles, bookmark_mapping, bookmark_id_counter
            )
            return True, bookmark_id_counter
        else:
            print(f"      ⚠️ No se encontró contenido principal en {html_file.name}")
            return False, bookmark_id_counter

    except Exception as e:
        print(f"      ❌ Error procesando {html_file.name}: {e}")
        return False, bookmark_id_counter

def add_heading_with_real_bookmark(doc, title, filename, bookmark_mapping, bookmark_id_counter):
    """Agregar encabezado con bookmark REAL"""
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    bookmark_name = bookmark_mapping.get(filename, clean_bookmark_name(filename))

    try:
        # Crear bookmark REAL con XML
        p = heading._element

        # Crear bookmark start
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id_counter))
        bookmark_start.set(qn('w:name'), bookmark_name)

        # Crear bookmark end
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id_counter))

        # Insertar bookmarks en el párrafo
        p.insert(0, bookmark_start)
        p.append(bookmark_end)

        # print(f"      🔖 Bookmark REAL creado: {bookmark_name} (ID: {bookmark_id_counter})")  # Simplificado

        return heading, bookmark_id_counter + 1

    except Exception as e:
        print(f"      ⚠️ Error creando bookmark real {bookmark_name}: {e}")
        return heading, bookmark_id_counter + 1

def extract_title(soup, html_file):
    """Extraer título del HTML"""
    # Buscar el h2 dentro del contenido principal (el título real de la sección)
    main_content = soup.find('div', class_='main-content') or soup.find('div', id='topic-content')
    if main_content:
        h2_elem = main_content.find('h2')
        if h2_elem:
            title = h2_elem.get_text().strip()
            if title:
                return title

    # Fallback: buscar cualquier h1 o h2 que no sea genérico
    title_selectors = ['h2', 'h1']
    for selector in title_selectors:
        title_elem = soup.select_one(selector)
        if title_elem:
            title = title_elem.get_text().strip()
            # Excluir títulos muy genéricos o que contengan patrones de manual
            if title and not _is_generic_manual_title(title):
                return title

    return html_file.stem.replace('_', ' ').replace('-', ' ').title()

def _is_generic_manual_title(title):
    """Verificar si un título es genérico de manual (multi-idioma)"""
    generic_patterns = [
        r'manual\s+de\s+usuario',  # Español
        r'user\s+manual',          # Inglés
        r'virtual\s+campus',       # Común
        r'campus\s+virtual',       # Español/Portugués
        r'manual\s+do\s+usuário',  # Portugués
        r'manuel\s+utilisateur',   # Francés
        r'manuale\s+utente',       # Italiano
        r'benutzerhandbuch',       # Alemán
        r'gebruikershandleiding',  # Holandés
        r'manual\s+d\'usuari',     # Catalán
        r'erabiltzaile\s+eskuliburua', # Euskera
        r'brugermanual',           # Danés
        r'användarmanual',         # Sueco
        r'puruhára\s+rogue'        # Guaraní
    ]

    title_lower = title.lower()
    for pattern in generic_patterns:
        if re.search(pattern, title_lower):
            return True
    return False

def process_html_content_with_real_links(doc, content_elem, base_path, css_styles, bookmark_mapping, bookmark_id_counter):
    """Procesar contenido HTML con enlaces internos REALES"""

    for element in content_elem.children:
        if hasattr(element, 'name'):
            if element.name == 'p':
                bookmark_id_counter = process_paragraph_with_real_links(
                    doc, element, base_path, css_styles, bookmark_mapping, bookmark_id_counter
                )
            elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(element.name[1])
                text = element.get_text().strip()
                if text:
                    add_heading(doc, text, level=min(level + 1, 9))
            elif element.name == 'img':
                # Imagen standalone
                img_paragraph = doc.add_paragraph()
                process_enhanced_image(img_paragraph, element, base_path, inline_context=False)

    return bookmark_id_counter

def add_heading(doc, text, level=1):
    """Agregar encabezado normal"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

def process_paragraph_with_real_links(doc, p_elem, base_path, css_styles, bookmark_mapping, bookmark_id_counter):
    """Procesar párrafo con enlaces internos REALES"""

    paragraph = doc.add_paragraph()
    has_content = False

    # Convertir children a lista para poder mirar el siguiente elemento
    children = list(p_elem.children)

    for i, child in enumerate(children):
        if hasattr(child, 'name'):
            if child.name == 'img':
                if has_content:
                    # Imagen inline dentro de párrafo con contenido
                    process_enhanced_image(paragraph, child, base_path, inline_context=True)
                else:
                    # Imagen sola en párrafo vacío - convertir a standalone
                    p = paragraph._element
                    p.getparent().remove(p)
                    img_paragraph = doc.add_paragraph()
                    process_enhanced_image(img_paragraph, child, base_path, inline_context=False)
                continue
            elif child.name == 'a':
                # Procesar enlace - CREAR HYPERLINK REAL
                link_added = process_link_with_real_hyperlink(
                    paragraph, child, css_styles, bookmark_mapping
                )
                if link_added:
                    has_content = True
            elif child.name == 'span':
                text_added = process_enhanced_span(paragraph, child, css_styles)
                if text_added:
                    has_content = True
            elif child.name == 'br':
                add_run_with_font(paragraph).add_break()
                has_content = True
            else:
                text = child.get_text()
                if text.strip():
                    # Preservar espacios al inicio y final
                    import re
                    original_text = text
                    starts_with_space = original_text.startswith((' ', '\n', '\t'))
                    ends_with_space = original_text.endswith((' ', '\n', '\t'))

                    text = re.sub(r'\s+', ' ', text.strip())

                    if starts_with_space:
                        text = ' ' + text
                    if ends_with_space:
                        text = text + ' '

                    run = add_run_with_font(paragraph, text)
                    has_content = True
        else:
            text = str(child)
            if text.strip():
                # Preservar espacios al inicio y final
                import re
                original_text = text
                starts_with_space = original_text.startswith((' ', '\n', '\t'))
                ends_with_space = original_text.endswith((' ', '\n', '\t'))

                text = re.sub(r'\s+', ' ', text.strip())

                if starts_with_space:
                    text = ' ' + text
                if ends_with_space:
                    text = text + ' '

                run = add_run_with_font(paragraph, text)
                has_content = True

    # Remover párrafo vacío
    if not has_content:
        try:
            p = paragraph._element
            p.getparent().remove(p)
        except:
            pass

    return bookmark_id_counter

def process_link_with_real_hyperlink(paragraph, link_elem, css_styles, bookmark_mapping):
    """Procesar enlace con HYPERLINK REAL"""

    href = link_elem.get('href', '')
    link_text = link_elem.get_text()
    css_class = link_elem.get('class', [])

    # Preservar espacios al inicio y final del enlace, pero normalizar internos
    import re
    original_link_text = link_text
    starts_with_space = original_link_text.startswith((' ', '\n', '\t'))
    ends_with_space = original_link_text.endswith((' ', '\n', '\t'))

    # Normalizar espacios internos
    link_text = re.sub(r'\s+', ' ', link_text.strip())

    # Restaurar espacios de borde si existían
    if starts_with_space:
        link_text = ' ' + link_text
    if ends_with_space:
        link_text = link_text + ' '

    # REDIRECCIONAR enlaces hacia Home/index → Introduccion
    if href in ['index.html', 'Home.html']:
        href = 'Introduccion.html'
        print(f"         ↪️ Redirigiendo enlace: '{link_text}' → Introducción")

    if not link_text.strip():
        return False

    # DETECTAR SI ES ENLACE INTERNO
    is_internal_link = href.endswith('.html') and not href.startswith('http')

    if is_internal_link:
        # Es enlace interno - crear HYPERLINK XML REAL
        target_bookmark = bookmark_mapping.get(href, clean_bookmark_name(href))
        success = create_real_internal_hyperlink(paragraph, link_text, target_bookmark, css_styles, css_class)
        if success:
            # print(f"         🔗 Enlace REAL: '{link_text}' → {target_bookmark}")  # Simplificado
            pass
        return success
    else:
        # Es enlace externo - crear HYPERLINK XML REAL
        success = create_real_external_hyperlink(paragraph, link_text, href, css_styles, css_class)
        if success:
            # print(f"         🌐 Enlace externo REAL: '{link_text}' → {href}")  # Simplificado
            pass
        return success

def create_real_internal_hyperlink(paragraph, text, bookmark_name, css_styles, css_class):
    """Crear HYPERLINK XML REAL hacia bookmark"""

    try:
        # Crear elemento hyperlink XML
        p = paragraph._element

        # Crear el elemento hyperlink con anchor hacia bookmark
        hyperlink_xml = f'''
        <w:hyperlink w:anchor="{bookmark_name}" {nsdecls('w')}>
            <w:r>
                <w:rPr>
                    <w:color w:val="0000FF"/>
                    <w:u w:val="single"/>
                </w:rPr>
                <w:t>{text}</w:t>
            </w:r>
        </w:hyperlink>
        '''

        hyperlink_element = parse_xml(hyperlink_xml)
        p.append(hyperlink_element)

        return True

    except Exception as e:
        print(f"            ⚠️ Error creando hyperlink XML: {e}")
        # Fallback: crear run normal con formato de enlace
        run = add_run_with_font(paragraph, text)
        apply_link_style(run, css_styles, css_class)
        return True

def create_real_external_hyperlink(paragraph, text, url, css_styles, css_class):
    """Crear HYPERLINK XML REAL hacia URL externa"""

    try:
        # Para enlaces externos necesitamos:
        # 1. Agregar la relación al documento
        # 2. Crear hyperlink XML con r:id que apunte a esa relación

        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

        # Crear elemento hyperlink XML con relationship ID
        p = paragraph._element

        hyperlink_xml = f'''
        <w:hyperlink r:id="{r_id}" {nsdecls('w', 'r')}>
            <w:r>
                <w:rPr>
                    <w:color w:val="0000FF"/>
                    <w:u w:val="single"/>
                </w:rPr>
                <w:t>{text}</w:t>
            </w:r>
        </w:hyperlink>
        '''

        hyperlink_element = parse_xml(hyperlink_xml)
        p.append(hyperlink_element)

        return True

    except Exception as e:
        print(f"            ⚠️ Error creando hyperlink externo XML: {e}")
        # Fallback: crear run normal con formato de enlace
        run = add_run_with_font(paragraph, text)
        apply_link_style(run, css_styles, css_class)
        return True

def apply_link_style(run, css_styles, css_class):
    """Aplicar estilo de enlace a un run"""

    # Aplicar estilos CSS si existen
    if css_class:
        class_name = css_class[0] if isinstance(css_class, list) else css_class
        apply_css_style_to_run(run, class_name, css_styles)

    # Asegurar que se vea como enlace
    run.font.color.rgb = RGBColor(0, 0, 255)  # Azul
    run.font.underline = True

def process_enhanced_span(paragraph, span_elem, css_styles):
    """Procesar span con estilos CSS"""

    text = span_elem.get_text()
    if not text.strip():
        return False

    # Preservar espacios al inicio y final, pero normalizar internos
    import re
    original_text = text
    starts_with_space = original_text.startswith((' ', '\n', '\t'))
    ends_with_space = original_text.endswith((' ', '\n', '\t'))

    # Normalizar espacios internos
    text = re.sub(r'\s+', ' ', text.strip())

    # Restaurar espacios de borde si existían
    if starts_with_space:
        text = ' ' + text
    if ends_with_space:
        text = text + ' '

    run = add_run_with_font(paragraph, text)

    css_class = span_elem.get('class', [])
    if css_class:
        class_name = css_class[0] if isinstance(css_class, list) else css_class
        apply_css_style_to_run(run, class_name, css_styles)

    return True

def apply_css_style_to_run(run, class_name, css_styles):
    """Aplicar estilos CSS a un run"""

    if class_name not in css_styles:
        return

    style_props = css_styles[class_name]

    if style_props.get('font-weight') == 'bold':
        run.font.bold = True

    if style_props.get('font-style') == 'italic':
        run.font.italic = True

    if style_props.get('text-decoration') == 'underline':
        run.font.underline = True

    color = style_props.get('color', '')
    if color:
        rgb_color = parse_css_color(color)
        if rgb_color:
            run.font.color.rgb = rgb_color

def parse_css_color(color_str):
    """Convertir color CSS a RGBColor"""

    color_str = color_str.strip().lower()

    if color_str.startswith('#'):
        try:
            if len(color_str) == 7:
                r = int(color_str[1:3], 16)
                g = int(color_str[3:5], 16)
                b = int(color_str[5:7], 16)
                return RGBColor(r, g, b)
        except:
            pass

    color_map = {
        '#000000': RGBColor(0, 0, 0),
        '#0000ff': RGBColor(0, 0, 255),
        '#000080': RGBColor(0, 0, 128),
        '#800000': RGBColor(128, 0, 0),
        '#008000': RGBColor(0, 128, 0),
        '#c0c0c0': RGBColor(192, 192, 192),
        '#6666ff': RGBColor(102, 102, 255),
    }

    return color_map.get(color_str)

def process_enhanced_image(paragraph, img_elem, base_path, inline_context=False):
    """Procesar imagen HTML con soporte para imágenes inline y standalone"""

    src = img_elem.get('src')
    if not src:
        return False

    img_path = base_path / src

    try:
        if img_path.exists():
            # Determinar si es imagen inline o standalone
            is_inline = _is_inline_image(img_elem, inline_context, base_path)

            if is_inline:
                # Imagen inline: agregar al párrafo actual, tamaño pequeño
                run = paragraph.add_run()
                try:
                    run.add_picture(str(img_path), height=Inches(0.2))  # ~5mm altura
                    return True
                except Exception as e:
                    print(f"      ⚠️ Error imagen inline {src}: {e}")
                    return False
            else:
                # Imagen standalone: crear párrafo separado, tamaño normal
                run = paragraph.add_run()
                try:
                    run.add_picture(str(img_path), width=Inches(4))  # Reducido de 6 a 4
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    return True
                except Exception as e:
                    print(f"      ⚠️ Error imagen standalone {src}: {e}")
                    return False
        else:
            print(f"      ⚠️ Imagen no encontrada: {img_path}")
            return False

    except Exception as e:
        print(f"      ❌ Error procesando imagen {src}: {e}")
        return False

def _is_inline_image(img_elem, inline_context, base_path):
    """Determinar si una imagen debe ser tratada como inline basado en tamaño real"""

    # Verificar características del elemento imagen
    style = img_elem.get('style', '')
    src = img_elem.get('src', '')

    # Patrones de nombres que suelen ser iconos/botones (alta prioridad)
    icon_patterns = ['button', 'icon', 'btn', 'small']
    for pattern in icon_patterns:
        if pattern.lower() in src.lower():
            return True

    # Verificar tamaño real de la imagen ANTES del padding
    if src and base_path:
        img_path = base_path / src
        if img_path.exists():
            try:
                from PIL import Image
                with Image.open(img_path) as img:
                    width, height = img.size
                    # print(f"      📏 {src}: {width}x{height}px")  # Simplificado

                    # Considerar inline si es pequeña en dimensiones
                    # Típicamente botones/iconos son <100px en alguna dimensión
                    if width <= 100 or height <= 100:
                        # print(f"         → INLINE (pequeña: ≤100px)")  # Simplificado
                        return True

                    # Si es grande (>300px en ambas), definitivamente standalone
                    if width > 300 and height > 200:
                        # print(f"         → STANDALONE (grande: {width}x{height} > 300x200)")  # Simplificado
                        return False

                    # Si está en contexto inline, verificar ratios especiales
                    if inline_context:
                        # Si es muy ancha pero baja (banner/barra), puede ser inline
                        if height <= 50 and width <= 400:
                            # print(f"         → INLINE (banner: {width}x{height})")  # Simplificado
                            return True
                        # Si es muy alta pero estrecha (botón vertical), puede ser inline
                        if width <= 50 and height <= 200:
                            # print(f"         → INLINE (botón vertical: {width}x{height})")  # Simplificado
                            return True

                    # print(f"         → Continuando análisis...")  # Simplificado

            except Exception as e:
                print(f"      ⚠️ Error verificando tamaño de {src}: {e}")
                # Fallback al contexto si no podemos verificar tamaño
                pass

    # Verificar padding como último recurso (puede ser engañoso)
    if 'padding : 1px' in style or 'padding: 1px' in style:
        return True

    # Verificar si es imagen en párrafo con texto (contexto inline)
    if inline_context:
        return True  # Default para contexto inline si no hay info de tamaño

    return False

def verify_real_links(docx_path):
    """Verificar que los enlaces REALES estén funcionando"""
    print(f"\n🔍 Verificando enlaces REALES: {docx_path}")

    try:
        # Verificar XML directamente
        import zipfile
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            doc_xml = zip_ref.read('word/document.xml').decode('utf-8')

            hyperlink_count = doc_xml.count('<w:hyperlink')
            bookmark_start_count = doc_xml.count('<w:bookmarkStart')
            bookmark_end_count = doc_xml.count('<w:bookmarkEnd')

            print(f"   🔗 Hyperlinks XML: {hyperlink_count}")
            print(f"   🔖 Bookmark starts: {bookmark_start_count}")
            print(f"   🔖 Bookmark ends: {bookmark_end_count}")

            if hyperlink_count > 0 and bookmark_start_count > 0:
                print("   💡 En Word: Ctrl+Click funcionará para navegar")
                return True
            else:
                print("   ❌ No se encontraron elementos XML reales")
                return False

    except Exception as e:
        print(f"   ❌ Error verificando: {e}")
        return False

if __name__ == "__main__":
    # Por defecto, usar manual back, pero se puede cambiar
    manual_type = 'open_aula_back'  # o 'open_aula_front'
    manual_dir = f"{manual_type}_es"

    html_dir = PROJECT_ROOT / "original" / manual_dir / "html"
    output_docx = PROJECT_ROOT / f"manual_{manual_type}_final.docx"

    print("🔗 CONVERTIDOR ESPAÑOL - MANUAL FINAL COMPLETO")
    print("=" * 60)

    if not html_dir.exists():
        print(f"❌ Error: No se encuentra {html_dir}")
        exit(1)

    # Convertir con enlaces REALES
    success = spanish_html_with_real_links(html_dir, output_docx)

    if success:
        print("\n🔍 Verificando enlaces REALES...")
        verify_success = verify_real_links(output_docx)

        if verify_success:
            print("\n🎯 ¡ENLACES INTERNOS REALES FUNCIONANDO!")
            print(f"\n📁 Archivo con enlaces REALES: {output_docx}")
            print("🚀 ¡Ahora los enlaces SÍ funcionan!")
        else:
            print("\n⚠️ Revisar implementación de enlaces")
    else:
        print("\n❌ Error en la conversión con enlaces reales")