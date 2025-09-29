#!/usr/bin/env python3
"""
TOC Handler - Manejo centralizado de estructura y generación de índices
Parseador de _toc.json para generar índices jerárquicos multi-idioma
"""

import json
import re
from pathlib import Path
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor


class TOCHandler:
    """Manejador centralizado de TOC para múltiples idiomas"""

    def __init__(self, manual_name='open_aula_front'):
        self.manual_name = manual_name

    def parse_toc_json(self, html_dir_path):
        """
        Parsear _toc.json para obtener estructura jerárquica

        Args:
            html_dir_path: Path al directorio con archivos HTML y _toc.json

        Returns:
            list: Estructura jerárquica del TOC
        """
        toc_file = html_dir_path / "_toc.json"

        if not toc_file.exists():
            print(f"   ⚠️ No se encontró _toc.json en {html_dir_path}")
            return []

        try:
            with open(toc_file, 'r', encoding='utf-8') as f:
                toc_data = json.load(f)

            print(f"   📋 TOC parseado: {len(toc_data)} elementos")

            # Convertir estructura plana a jerárquica
            return self._build_hierarchical_structure(toc_data)

        except Exception as e:
            print(f"   ❌ Error parseando _toc.json: {e}")
            return []

    def _build_hierarchical_structure(self, toc_data):
        """
        Convertir estructura plana de JSTree a jerárquica

        Args:
            toc_data: Lista de elementos del TOC desde _toc.json

        Returns:
            list: Estructura jerárquica
        """
        # Crear mapeo de elementos por ID
        elements = {item['id']: item for item in toc_data}

        # Crear estructura jerárquica
        root_items = []

        for item in toc_data:
            parent_id = item.get('parent', '#')

            # Agregar campos necesarios
            item['level'] = 0
            item['children'] = []
            item['href'] = item.get('a_attr', {}).get('href', '')
            item['title'] = item.get('text', '')

            # Skip items sin href válido o especiales
            if not item['href'] or item['href'].startswith('../') or item['id'] == 'Home':
                continue

            if parent_id == '#':
                # Elemento raíz
                root_items.append(item)
            else:
                # Elemento hijo - agregarlo al padre
                if parent_id in elements:
                    parent = elements[parent_id]
                    if 'children' not in parent:
                        parent['children'] = []
                    parent['children'].append(item)

        # Calcular niveles recursivamente
        self._calculate_levels(root_items, 0)

        print(f"   🌳 Estructura jerárquica: {len(root_items)} elementos raíz")
        return root_items

    def _calculate_levels(self, items, level):
        """Calcular niveles recursivamente"""
        for item in items:
            item['level'] = level
            if item.get('children'):
                self._calculate_levels(item['children'], level + 1)

    def get_ordered_html_files(self, html_dir_path, toc_structure):
        """
        Obtener archivos HTML ordenados según estructura TOC

        Args:
            html_dir_path: Path al directorio HTML
            toc_structure: Estructura jerárquica del TOC

        Returns:
            list: Lista de Path objects ordenados según TOC
        """
        # Crear lista plana de archivos en orden TOC
        ordered_hrefs = []
        self._collect_hrefs_recursive(toc_structure, ordered_hrefs)

        # Mapear a archivos reales
        available_files = {f.name: f for f in html_dir_path.glob('*.html')}
        ordered_files = []

        for href in ordered_hrefs:
            if href in available_files:
                ordered_files.append(available_files[href])

        # Agregar archivos que no están en TOC al final
        for filename, filepath in available_files.items():
            if filename not in ordered_hrefs and filename not in ['index.html', 'Home.html', 'DescargarmanualenPDF.html']:
                ordered_files.append(filepath)

        print(f"   📊 Archivos ordenados: {len(ordered_files)}")
        return ordered_files

    def _collect_hrefs_recursive(self, items, href_list):
        """Recopilar hrefs recursivamente en orden"""
        for item in items:
            if item.get('href'):
                href_list.append(item['href'])
            if item.get('children'):
                self._collect_hrefs_recursive(item['children'], href_list)

    def create_hierarchical_index(self, doc, toc_structure, available_files, lang_code='en'):
        """
        Crear índice jerárquico con numeración automática

        Args:
            doc: Documento DOCX
            toc_structure: Estructura jerárquica del TOC
            available_files: Dict de archivos disponibles {nombre: Path}
            lang_code: Código de idioma
        """
        from html_to_docx import add_run_with_font, clean_bookmark_name, create_real_internal_hyperlink

        # Obtener traducciones para el título del índice
        index_titles = {
            'es': 'Índice',
            'en': 'Index',
            'pt': 'Índice',
            'fr': 'Index',
            'it': 'Indice',
            'de': 'Inhaltsverzeichnis',
            'nl': 'Inhoudsopgave',
            'ca': 'Índex',
            'eu': 'Aurkibidea',
            'gl': 'Índice',
            'da': 'Indeks',
            'sv': 'Index',
            'gn': 'Jehaipyre'
        }

        index_title = index_titles.get(lang_code, 'Index')

        # Título del índice
        index_heading = doc.add_paragraph()
        index_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = add_run_with_font(index_heading, index_title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # Espacio
        doc.add_paragraph()

        print(f"   📋 Generando índice jerárquico para {lang_code}...")

        # Generar entradas del índice
        self._add_toc_items_recursive(doc, toc_structure, available_files, [])

        # Salto de página después del índice
        doc.add_page_break()

    def _add_toc_items_recursive(self, doc, items, available_files, numbering_path):
        """Agregar elementos del TOC recursivamente con numeración"""
        from html_to_docx import add_run_with_font, clean_bookmark_name, create_real_internal_hyperlink

        for i, item in enumerate(items, 1):
            href = item.get('href', '')
            title = item.get('title', '')
            level = item.get('level', 0)

            if href in available_files:
                # Calcular numeración jerárquica
                current_numbering = numbering_path + [i]
                number_str = '.'.join(map(str, current_numbering)) + '.'

                # Crear entrada de índice
                index_para = doc.add_paragraph()

                # Indentación basada en nivel
                indent = Pt(36 + (level * 18))  # 36pt base + 18pt por nivel
                index_para.paragraph_format.left_indent = indent

                # Agregar numeración
                num_run = add_run_with_font(index_para, f"{number_str} ")
                num_run.font.bold = True
                num_run.font.color.rgb = RGBColor(0, 51, 102)

                # Crear hyperlink para el título
                bookmark_name = clean_bookmark_name(href)
                create_real_internal_hyperlink(index_para, title, bookmark_name, {}, [])

                # print(f"     📋 Índice: {'  ' * level}{number_str} {title} → {bookmark_name}")  # Simplificado

            # Procesar hijos recursivamente
            if item.get('children'):
                child_numbering = numbering_path + [i] if href in available_files else numbering_path
                self._add_toc_items_recursive(doc, item['children'], available_files, child_numbering)


def clean_bookmark_name(filename):
    """Limpiar nombre de archivo para usar como bookmark"""
    name = filename.replace('.html', '')
    name = re.sub(r'[^a-zA-Z0-9]', '_', name)
    if name and name[0].isdigit():
        name = 'section_' + name
    return name or 'unknown'