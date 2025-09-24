#!/usr/bin/env python3
"""
Conversor HTML a DOCX Multi-idioma
Versión adaptada de spanish_converter_real_links.py para soporte multi-idioma
"""

import sys
from pathlib import Path

# Agregar directorios al path
sys.path.append(str(Path(__file__).parent))

from languages_config import LANGUAGES
from system_config import DOCX_CONFIG, get_manual_path
from html_to_docx import (
    setup_enhanced_document, load_css_styles_from_spanish, get_all_html_files_structured,
    create_bookmark_mapping, process_html_file_with_real_links, create_title_page_and_index,
    DOCXLogger, DOCXProgressDisplay
)

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

class MultiLanguageDocxConverter:
    """Conversor HTML a DOCX que soporta múltiples idiomas"""

    def __init__(self, manual_name='open_aula_front'):
        self.manual_name = manual_name

    def get_title_page_text(self, lang_code):
        """Obtiene el texto de la página de título según el idioma y tipo de manual"""
        # Determinar si es manual de usuario o administración
        is_admin_manual = 'back' in self.manual_name

        if lang_code in DOCX_CONFIG['title_page_title']:
            base_title = DOCX_CONFIG['title_page_title'][lang_code]
        else:
            # Para idiomas no configurados, usar traducciones básicas
            if is_admin_manual:
                # Traducciones para Manual de Administración
                translations = {
                    'en': 'Virtual Campus\nAdministration Manual',
                    'pt': 'Campus Virtual\nManual de Administração',
                    'fr': 'Campus Virtuel\nManuel d\'Administration',
                    'it': 'Campus Virtuale\nManuale di Amministrazione',
                    'de': 'Virtueller Campus\nVerwaltungshandbuch',
                    'nl': 'Virtuele Campus\nBeheerhandleiding',
                    'ca': 'Campus Virtual\nManual d\'Administració',
                    'eu': 'Campus Birtuala\nKudeaketa Eskuliburua',
                    'gl': 'Campus Virtual\nManual de Administración',
                    'da': 'Virtuelt Campus\nAdministrationsmanual',
                    'sv': 'Virtuellt Campus\nAdministrationsmanual',
                    'gn': 'Campus Virtual\nÑangarekoha Rogue'
                }
            else:
                # Traducciones para Manual de Usuario
                translations = {
                    'en': 'Virtual Campus\nUser Manual',
                    'pt': 'Campus Virtual\nManual do Usuário',
                    'fr': 'Campus Virtuel\nManuel Utilisateur',
                    'it': 'Campus Virtuale\nManuale Utente',
                    'de': 'Virtueller Campus\nBenutzerhandbuch',
                    'nl': 'Virtuele Campus\nGebruikershandleiding',
                    'ca': 'Campus Virtual\nManual d\'Usuari',
                    'eu': 'Campus Birtuala\nErabiltzaile Eskuliburua',
                    'gl': 'Campus Virtual\nManual de Usuario',
                    'da': 'Virtuelt Campus\nBrugermanual',
                    'sv': 'Virtuellt Campus\nAnvändarmanual',
                    'gn': 'Campus Virtual\nPuruhára Rogue'
                }

            base_title = translations.get(lang_code, DOCX_CONFIG['title_page_title']['default'])

        return base_title

    def get_footer_translations(self, lang_code):
        """Obtener traducciones para el footer según idioma"""
        from datetime import datetime
        current_year = datetime.now().year

        translations = {
            'es': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Página',
                'of': 'de'
            },
            'en': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Page',
                'of': 'of'
            },
            'pt': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Página',
                'of': 'de'
            },
            'fr': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Page',
                'of': 'de'
            },
            'it': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Pagina',
                'of': 'di'
            },
            'de': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Seite',
                'of': 'von'
            },
            'nl': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Pagina',
                'of': 'van'
            },
            'ca': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Pàgina',
                'of': 'de'
            },
            'eu': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Orria',
                'of': 'eko'
            },
            'gl': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Páxina',
                'of': 'de'
            },
            'da': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Side',
                'of': 'af'
            },
            'sv': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Sida',
                'of': 'av'
            },
            'gn': {
                'company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'page': 'Kuatiarogue',
                'of': 'peve'
            }
        }

        return translations.get(lang_code, translations['es'])

    def patch_translations_for_docx(self, lang_code):
        """Parche temporal para agregar traducciones completas a la biblioteca html_to_docx"""
        import html_to_docx
        from datetime import datetime
        current_year = datetime.now().year

        # Determinar si es manual de usuario o administración
        is_admin_manual = 'back' in self.manual_name

        # Traducciones completas para todos los idiomas
        all_translations = {
            'es': {
                'title': 'Campus Virtual' if not is_admin_manual else 'Campus Virtual',
                'subtitle': 'Manual de Usuario' if not is_admin_manual else 'Manual de Administración',
                'index_title': 'Índice',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Página',
                'footer_of': 'de'
            },
            'en': {
                'title': 'Virtual Campus',
                'subtitle': 'User Manual' if not is_admin_manual else 'Administration Manual',
                'index_title': 'Index',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Page',
                'footer_of': 'of'
            },
            'pt': {
                'title': 'Campus Virtual',
                'subtitle': 'Manual do Usuário' if not is_admin_manual else 'Manual de Administração',
                'index_title': 'Índice',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Página',
                'footer_of': 'de'
            },
            'fr': {
                'title': 'Campus Virtuel',
                'subtitle': 'Manuel Utilisateur' if not is_admin_manual else 'Manuel d\'Administration',
                'index_title': 'Index',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Page',
                'footer_of': 'de'
            },
            'it': {
                'title': 'Campus Virtuale',
                'subtitle': 'Manuale Utente' if not is_admin_manual else 'Manuale di Amministrazione',
                'index_title': 'Indice',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Pagina',
                'footer_of': 'di'
            },
            'de': {
                'title': 'Virtueller Campus',
                'subtitle': 'Benutzerhandbuch' if not is_admin_manual else 'Verwaltungshandbuch',
                'index_title': 'Inhalt',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Seite',
                'footer_of': 'von'
            },
            'nl': {
                'title': 'Virtuele Campus',
                'subtitle': 'Gebruikershandleiding' if not is_admin_manual else 'Beheerhandleiding',
                'index_title': 'Inhoudsopgave',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Pagina',
                'footer_of': 'van'
            },
            'ca': {
                'title': 'Campus Virtual',
                'subtitle': 'Manual d\'Usuari' if not is_admin_manual else 'Manual d\'Administració',
                'index_title': 'Índex',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Pàgina',
                'footer_of': 'de'
            },
            'eu': {
                'title': 'Campus Birtuala',
                'subtitle': 'Erabiltzaile Eskuliburua' if not is_admin_manual else 'Kudeaketa Eskuliburua',
                'index_title': 'Aurkibidea',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Orria',
                'footer_of': 'eko'
            },
            'gl': {
                'title': 'Campus Virtual',
                'subtitle': 'Manual de Usuario' if not is_admin_manual else 'Manual de Administración',
                'index_title': 'Índice',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Páxina',
                'footer_of': 'de'
            },
            'da': {
                'title': 'Virtuelt Campus',
                'subtitle': 'Brugermanual' if not is_admin_manual else 'Administrationsmanual',
                'index_title': 'Indholdsfortegnelse',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Side',
                'footer_of': 'af'
            },
            'sv': {
                'title': 'Virtuellt Campus',
                'subtitle': 'Användarmanual' if not is_admin_manual else 'Administrationsmanual',
                'index_title': 'Innehållsförteckning',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Sida',
                'footer_of': 'av'
            },
            'gn': {
                'title': 'Campus Virtual',
                'subtitle': 'Puruhára Rogue' if not is_admin_manual else 'Ñangarekoha Rogue',
                'index_title': 'Índice',
                'footer_company': f'© {current_year} e-ducativa Educación Virtual S.A.',
                'footer_page': 'Kuatiarogue',
                'footer_of': 'peve'
            }
        }

        # Función de reemplazo para get_translations
        def patched_get_translations(language='es'):
            return all_translations.get(language, all_translations['es'])

        # Aplicar el parche
        html_to_docx.get_translations = patched_get_translations
        print(f"   🔧 Parche de traducciones aplicado para {lang_code}")

    def convert_docx_to_pdf(self, docx_path):
        """
        VERSIÓN MEJORADA: Convierte DOCX a PDF usando wkhtmltopdf patched con correcciones
        """
        import tempfile

        try:
            # Crear directorio pdf/
            docx_dir = docx_path.parent
            pdf_dir = docx_dir.parent / 'pdf'
            pdf_dir.mkdir(parents=True, exist_ok=True)

            pdf_filename = docx_path.stem + '.pdf'
            pdf_path = pdf_dir / pdf_filename

            print(f"   📄 Convirtiendo a PDF mejorado: {pdf_path.name}")

            # Obtener idioma del archivo
            lang_parts = pdf_filename.split('_')
            current_lang = lang_parts[-1].replace('.pdf', '') if len(lang_parts) > 1 else 'es'

            # Crear directorio temporal
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)

                # Paso 1: DOCX → HTML
                html_file = temp_path / 'document.html'
                print("   1️⃣ Convirtiendo DOCX a HTML...")

                result = subprocess.run([
                    'pandoc', str(docx_path), '-o', str(html_file),
                    '--extract-media', str(temp_path), '--standalone'
                ], capture_output=True, text=True)

                if result.returncode != 0:
                    return False, None, f"Error pandoc: {result.stderr}"

                # Paso 2: Limpiar HTML
                print("   2️⃣ Limpiando HTML...")
                with open(html_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()

                # Arreglar rutas de imágenes
                html_content = html_content.replace('media/media/', 'media/')

                # Limpiar títulos duplicados y marcar elementos especiales
                lines = html_content.split('\n')
                cleaned_lines = []
                seen_title = False
                seen_subtitle = False

                # Detectar el tipo de manual para usar las traducciones correctas
                is_admin_manual = 'back' in str(docx_path) or 'administracion' in str(docx_path).lower()

                # Traducciones hardcodeadas para Campus Virtual
                campus_translations = {
                    'it': 'Campus Virtuale',
                    'es': 'Campus Virtual',
                    'en': 'Virtual Campus',
                    'pt': 'Campus Virtual',
                    'fr': 'Campus Virtuel',
                    'de': 'Virtueller Campus',
                    'nl': 'Virtuele Campus',
                    'ca': 'Campus Virtual',
                    'eu': 'Kanpus Birtuala',
                    'gl': 'Campus Virtual',
                    'da': 'Virtuelt Campus',
                    'sv': 'Virtuell Campus',
                    'gn': 'Tetã Arandukarenda'
                }

                # Traducciones hardcodeadas para Manual de Usuario/Administración
                if is_admin_manual:
                    manual_translations = {
                        'it': 'Manuale di Amministrazione',
                        'es': 'Manual de Administración',
                        'en': 'Administration Manual',
                        'pt': 'Manual de Administração',
                        'fr': 'Manuel d\'Administration',
                        'de': 'Administrationshandbuch',
                        'nl': 'Administratiehandleiding',
                        'ca': 'Manual d\'Administració',
                        'eu': 'Administrazio Eskuliburua',
                        'gl': 'Manual de Administración',
                        'da': 'Administrationshåndbog',
                        'sv': 'Administrationshandbok',
                        'gn': 'Ñangareko Kuatiañe\'ẽ'
                    }
                else:
                    manual_translations = {
                        'it': 'Manuale Utente',
                        'es': 'Manual de Usuario',
                        'en': 'User Manual',
                        'pt': 'Manual do Usuário',
                        'fr': 'Manuel Utilisateur',
                        'de': 'Benutzerhandbuch',
                        'nl': 'Gebruikershandleiding',
                        'ca': 'Manual d\'Usuari',
                        'eu': 'Erabiltzaile Eskuliburua',
                        'gl': 'Manual de Usuario',
                        'da': 'Brugerhåndbog',
                        'sv': 'Användarhandbok',
                        'gn': 'Puruhára Kuatiañe\'ẽ'
                    }

                # Traducciones hardcodeadas para Índice
                index_translations = {
                    'it': 'Indice',
                    'es': 'Índice',
                    'en': 'Index',
                    'pt': 'Índice',
                    'fr': 'Index',
                    'de': 'Inhaltsverzeichnis',
                    'nl': 'Inhoudsopgave',
                    'ca': 'Índex',
                    'eu': 'Aurkibidea',
                    'gl': 'Índice',
                    'da': 'Indholdsfortegnelse',
                    'sv': 'Innehållsförteckning',
                    'gn': 'Kuatia Ñe\'ẽrã'
                }

                # Obtener las traducciones para el idioma actual
                campus_text = campus_translations.get(current_lang, 'Campus Virtual')
                manual_text = manual_translations.get(current_lang, 'Manual de Usuario')
                index_text = index_translations.get(current_lang, 'Índice')

                # Detectar títulos por POSICIÓN y ESTRUCTURA, reemplazar con traducciones hardcodeadas
                title_count = 0

                for line in lines:
                    # Detectar elementos que tienen contenido real (no vacío, no solo espacios)
                    has_content = False
                    if '<h1>' in line or '<h2>' in line or '<p><strong>' in line:
                        # Extraer el texto del elemento
                        import re
                        text_match = re.search(r'>([^<]+)<', line)
                        if text_match and text_match.group(1).strip() and len(text_match.group(1).strip()) > 2:
                            has_content = True
                            title_count += 1

                    # Los primeros 3 elementos con contenido son los títulos principales
                    if has_content and title_count == 1 and not seen_title:
                        # Primer título = Campus Virtual (traducido y con clase especial)
                        cleaned_lines.append(f'<h1 class="main-title">{campus_text}</h1>')
                        seen_title = True
                    # Segundo elemento con contenido = Manual de Usuario/Administración
                    elif has_content and title_count == 2 and not seen_subtitle:
                        # Segundo título = Manual de Usuario/Administración (traducido y con clase especial)
                        cleaned_lines.append(f'<h2 class="main-subtitle">{manual_text}</h2>')
                        seen_subtitle = True
                    # Tercer elemento con contenido = Índice
                    elif has_content and title_count == 3:
                        # Tercer título = Índice (traducido y con clase especial)
                        cleaned_lines.append(f'<h2 class="index-title">{index_text}</h2>')
                    # Omitir duplicados de los primeros 3 títulos principales después de procesarlos
                    elif has_content and title_count <= 3 and (seen_title and seen_subtitle):
                        continue  # Saltar duplicados de títulos principales
                    else:
                        cleaned_lines.append(line)

                html_content = '\n'.join(cleaned_lines)

                # Agregar CSS mejorado
                css = '''
<style>
@page {
    margin: 15mm 20mm 20mm 20mm;
}
body {
    font-family: Calibri, Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.3;
    margin: 0 auto;
    padding: 0;
    max-width: 90%;
    text-align: justify;
}

/* Título principal de la primera página - usar clase específica */
h1.main-title {
    font-size: 32pt !important;
    color: #17365D !important;
    text-align: center !important;
    margin: 60px 0 10px 0 !important;
    font-weight: bold !important;
    font-family: Calibri, Arial, sans-serif !important;
    page-break-after: avoid !important;
}

/* SIN línea decorativa */

/* Subtítulo principal (Manual de Usuario) */
h2.main-subtitle {
    font-size: 24pt !important;
    color: #4F81BD !important;
    text-align: center !important;
    margin: 0 0 40px 0 !important;
    font-weight: bold !important;
    font-family: Calibri, Arial, sans-serif !important;
    page-break-after: avoid !important;
}

/* Título del índice */
h2.index-title {
    font-size: 20pt !important;
    color: #365F91 !important;
    text-align: center !important;
    margin: 30px 0 20px 0 !important;
    font-weight: bold !important;
    page-break-before: avoid !important;
}

/* Encabezados normales de secciones - SIN línea decorativa */
h1:not(.main-title) {
    font-size: 14pt !important;
    color: #365F91 !important;
    font-weight: bold !important;
    margin-top: 24px !important;
    margin-bottom: 12px !important;
    page-break-before: always !important;
    text-align: left !important;
}

/* Asegurar que NO haya línea después de h1 normales */
h1:not(.main-title)::after {
    content: none !important;
}

h2:not(.main-subtitle):not(.index-title) {
    font-size: 13pt !important;
    color: #4F81BD !important;
    font-weight: bold !important;
    margin-top: 18px !important;
    margin-bottom: 10px !important;
    text-align: left !important;
}

h3 {
    font-size: 11pt !important;
    color: #4F81BD !important;
    font-weight: bold !important;
    margin-top: 12px !important;
    margin-bottom: 8px !important;
    text-align: left !important;
}

/* Enlaces */
a {
    color: #0563C1;
    text-decoration: underline;
}

/* Imágenes */
img {
    max-width: 100%;
    height: auto;
    display: block;
    margin: 10px auto;
}

/* Párrafos - ancho más conservador como Word */
p {
    margin: 8px 0;
    text-align: justify;
}

/* Tabla de contenidos */
.toc, ul.toc, ol.toc {
    margin-left: 20px;
    line-height: 1.8;
}

.toc a, ul.toc a, ol.toc a {
    text-decoration: none;
    color: #0563C1;
    font-size: 11pt;
}

/* Tablas - ancho más conservador */
table {
    border-collapse: collapse;
    width: 95%;
    margin: 15px auto;
}

td, th {
    border: 1px solid #ddd;
    padding: 8px;
    text-align: left;
}
</style>
'''
                html_content = html_content.replace('</head>', css + '\n</head>')

                # Guardar HTML limpio
                cleaned_html = temp_path / 'clean.html'
                with open(cleaned_html, 'w', encoding='utf-8') as f:
                    f.write(html_content)

                # Paso 3: PDF con wkhtmltopdf patched
                print("   3️⃣ Generando PDF...")
                wkhtmltopdf = Path('/home/sebos/tmp/traductor_manuales/bin/wkhtmltopdf_patched')

                if not wkhtmltopdf.exists():
                    return False, None, "wkhtmltopdf patched no encontrado"

                # Footer con formato igual al DOCX
                footer_translations = self.get_footer_translations(current_lang)

                result = subprocess.run([
                    str(wkhtmltopdf),
                    '--enable-local-file-access',
                    '--enable-internal-links',
                    '--outline',
                    '--outline-depth', '3',
                    '--margin-top', '15mm',
                    '--margin-bottom', '20mm',
                    '--margin-left', '15mm',
                    '--margin-right', '15mm',
                    '--footer-left', footer_translations['company'],
                    '--footer-right', f"{footer_translations['page']} [page] {footer_translations['of']} [topage]",
                    '--footer-font-size', '9',
                    '--footer-font-name', 'Calibri',
                    '--footer-spacing', '5',
                    str(cleaned_html),
                    str(pdf_path)
                ], capture_output=True, text=True, timeout=180)

                if result.returncode != 0:
                    return False, None, f"Error wkhtmltopdf: {result.stderr}"

                if pdf_path.exists():
                    file_size = pdf_path.stat().st_size
                    return True, pdf_path, f"PDF generado ({file_size:,} bytes)"
                else:
                    return False, None, "PDF no generado"

        except Exception as e:
            return False, None, f"Error: {str(e)}"

    def create_title_page_multilang(self, doc, lang_code):
        """Crear página de título adaptada al idioma"""
        title_text = self.get_title_page_text(lang_code)
        lines = title_text.split('\n')

        # Título principal (primera línea)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(lines[0])
        title_run.font.name = DOCX_CONFIG['default_font']
        title_run.font.size = Pt(24)
        title_run.font.bold = True

        # Subtítulo (segunda línea)
        if len(lines) > 1:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(lines[1])
            subtitle_run.font.name = DOCX_CONFIG['default_font']
            subtitle_run.font.size = Pt(18)
            subtitle_run.font.bold = True

        # Espacio
        doc.add_paragraph()
        doc.add_paragraph()

    def create_footer_multilang(self, doc, lang_code):
        """Crear pie de página con tabla invisible (formato profesional)"""
        from docx.oxml.shared import qn, OxmlElement
        from docx.shared import RGBColor

        print(f"   📄 Configurando pie de página para {lang_code}...")

        # Obtener traducciones de footer
        footer_translations = self.get_footer_translations(lang_code)

        section = doc.sections[0]
        footer = section.footer

        # Limpiar párrafos existentes
        for para in footer.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Crear tabla invisible de 1 fila x 2 columnas
        table = footer.add_table(rows=1, cols=2, width=Inches(7.5))
        table.style = 'Table Grid'

        # Configurar tabla invisible (sin bordes)
        for row in table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.tcPr
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.append(tcPr)

                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
                    border = OxmlElement(border_name)
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # Configurar ancho de columnas
        table.columns[0].width = Inches(4.5)  # Izquierda más ancha
        table.columns[1].width = Inches(3.0)  # Derecha más estrecha

        # CELDA IZQUIERDA: Copyright
        left_cell = table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run(footer_translations['company'])
        left_run.font.name = 'Calibri'
        left_run.font.size = Pt(9)
        left_run.font.color.rgb = RGBColor(102, 102, 102)

        # CELDA DERECHA: Numeración de páginas
        right_cell = table.cell(0, 1)
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Texto "Página " / "Page "
        page_run = right_para.add_run(f"{footer_translations['page']} ")
        page_run.font.name = 'Calibri'
        page_run.font.size = Pt(9)
        page_run.font.color.rgb = RGBColor(102, 102, 102)

        # Número de página actual
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        page_run._element.append(fldChar1)
        page_run._element.append(instrText)
        page_run._element.append(fldChar2)

        # " de " / " of "
        of_run = right_para.add_run(f" {footer_translations['of']} ")
        of_run.font.name = 'Calibri'
        of_run.font.size = Pt(9)
        of_run.font.color.rgb = RGBColor(102, 102, 102)

        # Número total de páginas
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        of_run._element.append(fldChar3)
        of_run._element.append(instrText2)
        of_run._element.append(fldChar4)

    def setup_document_multilang(self, doc, lang_code):
        """Configurar documento para idioma específico"""
        lang_info = LANGUAGES.get(lang_code, {})

        # Propiedades del documento
        core_props = doc.core_properties
        core_props.title = f"Manual de Usuario - {lang_info.get('name', lang_code)}"
        core_props.subject = "Manual de Usuario"
        core_props.author = "e-ducativa Educación Virtual S.A."
        core_props.comments = f"Manual traducido a {lang_info.get('name', lang_code)}"

        # Configurar fuente por defecto
        style = doc.styles['Normal']
        font = style.font
        font.name = DOCX_CONFIG['default_font']
        font.size = Pt(DOCX_CONFIG['styles']['normal']['size'])

        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def convert_html_to_docx(self, lang_code, force_regenerate=False):
        """
        Convierte HTML de un idioma específico a DOCX

        Args:
            lang_code: Código del idioma (ej: 'en', 'pt-BR')
            force_regenerate: Si True, regenera aunque ya exista

        Returns:
            tuple: (success: bool, output_path: Path, message: str)
        """

        # Validar idioma
        if lang_code == 'es':
            return False, None, "El español no se genera (viene del HelpNDoc original)"

        if lang_code not in LANGUAGES:
            return False, None, f"Idioma no soportado: {lang_code}"

        # Rutas de entrada y salida
        html_input_path = get_manual_path(self.manual_name, lang_code, 'html')
        docx_output_dir = get_manual_path(self.manual_name, lang_code, 'docx')
        docx_output_dir.mkdir(parents=True, exist_ok=True)

        # Determinar tipo de manual para el nombre del archivo
        if 'front' in self.manual_name:
            manual_type = 'aula_front'
        elif 'back' in self.manual_name:
            manual_type = 'aula_back'
        else:
            manual_type = 'aula_front'  # default

        output_file = docx_output_dir / f"manual_{manual_type}_{lang_code}.docx"

        # Verificar si HTML existe
        if not html_input_path.exists():
            return False, None, f"No existe HTML traducido para {LANGUAGES[lang_code]['name']}"

        html_files = list(html_input_path.glob('*.html'))
        if not html_files:
            return False, None, f"No se encontraron archivos HTML en {html_input_path}"

        # Verificar si ya existe DOCX
        if output_file.exists() and not force_regenerate:
            return True, output_file, f"DOCX ya existe para {LANGUAGES[lang_code]['name']}"

        try:
            # Inicializar logger y progreso
            logger = DOCXLogger(self.manual_name, lang_code)
            logger.log_step(f"DOCX_CONVERSION_START: {html_input_path} -> {output_file}")

            # Obtener archivos HTML primero para configurar progreso
            html_files_structured = get_all_html_files_structured(html_input_path)
            progress = DOCXProgressDisplay(len(html_files_structured))

            progress.show_start(output_file)
            logger.log_step(f"HTML_FILES_FOUND: {len(html_files_structured)} files")

            # Crear documento DOCX
            progress.show_step("Configurando documento")
            doc = Document()
            self.setup_document_multilang(doc, lang_code)
            logger.log_step("DOCUMENT_SETUP: Complete")

            # Cargar estilos CSS
            progress.show_step("Cargando estilos CSS")
            css_styles = load_css_styles_from_spanish()
            logger.log_step(f"CSS_STYLES_LOADED: {len(css_styles)} classes")

            # Aplicar parche de traducciones antes de crear la página de título
            progress.show_step("Aplicando parche de traducciones")
            self.patch_translations_for_docx(lang_code)
            logger.log_step("TRANSLATION_PATCH: Applied")

            # Crear página de título e índice
            progress.show_step("Creando índice")
            create_title_page_and_index(doc, html_files_structured, lang_code)
            logger.log_step("TITLE_PAGE_INDEX: Created")

            # Crear mapeo de bookmarks
            progress.show_step("Creando bookmarks")
            bookmark_mapping = create_bookmark_mapping(html_files_structured)
            logger.log_step(f"BOOKMARKS_CREATED: {len(bookmark_mapping)} bookmarks")

            # Procesar archivos HTML
            processed_files = 0
            bookmark_id_counter = 1

            for i, html_file in enumerate(html_files_structured):
                try:
                    progress.show_file_progress(i+1, len(html_files_structured), html_file.name)

                    success, bookmark_id_counter = process_html_file_with_real_links(
                        doc, html_file, html_input_path, css_styles, bookmark_mapping, bookmark_id_counter
                    )

                    if success:
                        processed_files += 1
                        logger.log_file_processed(html_file.name, True)
                    else:
                        logger.log_file_processed(html_file.name, False)

                    # Salto de página entre archivos
                    if i < len(html_files_structured) - 1:
                        doc.add_page_break()

                except Exception as e:
                    logger.log_step(f"ERROR_PROCESSING: {html_file.name} - {str(e)}")
                    continue

            # Configurar pie de página
            progress.show_step("Configurando pie de página")
            self.create_footer_multilang(doc, lang_code)
            logger.log_step("FOOTER: Created")

            # Guardar documento
            progress.show_step("Guardando documento")
            doc.save(str(output_file))

            # Obtener tamaño del archivo
            file_size = None
            if output_file.exists():
                file_size = output_file.stat().st_size

            logger.finalize(len(html_files_structured), processed_files, file_size)
            progress.show_complete(processed_files, len(html_files_structured), output_file)

            # Generar PDF automáticamente después del DOCX
            print(f"\n🔄 Generando PDF automáticamente...")
            pdf_success, pdf_path, pdf_message = self.convert_docx_to_pdf(output_file)

            if pdf_success:
                logger.log_step(f"PDF_GENERATED: {pdf_path}")
                return True, output_file, f"DOCX y PDF generados exitosamente para {LANGUAGES[lang_code]['name']}"
            else:
                logger.log_step(f"PDF_ERROR: {pdf_message}")
                return True, output_file, f"DOCX generado exitosamente para {LANGUAGES[lang_code]['name']}, pero error en PDF: {pdf_message}"

        except Exception as e:
            error_msg = f"Error generando DOCX para {LANGUAGES[lang_code]['name']}: {str(e)}"
            if 'logger' in locals():
                logger.log_step(f"DOCX_ERROR: {str(e)}")
            print(f"❌ {error_msg}")
            return False, None, error_msg

    def convert_multiple_languages(self, lang_codes, force_regenerate=False):
        """
        Convierte múltiples idiomas a DOCX

        Args:
            lang_codes: Lista de códigos de idioma
            force_regenerate: Si True, regenera aunque ya existan

        Returns:
            dict: Resultados por idioma
        """
        results = {}

        for lang_code in lang_codes:
            if lang_code == 'es':
                results[lang_code] = {
                    'success': False,
                    'message': 'Español no se genera (original HelpNDoc)'
                }
                continue

            success, output_path, message = self.convert_html_to_docx(lang_code, force_regenerate)
            results[lang_code] = {
                'success': success,
                'output_path': output_path,
                'message': message
            }

        return results

def main():
    """Función principal para usar desde línea de comandos"""
    import argparse

    parser = argparse.ArgumentParser(description='Convertir HTML a DOCX multi-idioma')
    parser.add_argument('--lang', required=True, help='Código de idioma (ej: en, pt, fr)')
    parser.add_argument('--manual', default='open_aula_front', help='Nombre del manual')
    parser.add_argument('--force', action='store_true', help='Forzar regeneración')

    args = parser.parse_args()

    converter = MultiLanguageDocxConverter(args.manual)
    success, output_path, message = converter.convert_html_to_docx(args.lang, args.force)

    if success:
        print(f"✅ {message}")
        if output_path:
            print(f"📄 Archivo: {output_path}")
        sys.exit(0)
    else:
        print(f"❌ {message}")
        sys.exit(1)

if __name__ == "__main__":
    main()